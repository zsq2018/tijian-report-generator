#!/usr/bin/env python3
"""
批量体检报告替换工具
用法: python3 batch_tihuan.py 人员信息模板.xlsx
"""
import json, sys, os, shutil, subprocess
from datetime import datetime, time

WORK_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE = os.path.join(WORK_DIR, "王显勇.docx")
SCRIPT = os.path.join(WORK_DIR, "tihuan_tijian.py")

COLUMN_MAP = {
    '姓名': 'name', '身份证号': 'idcard', '体检编码': 'number',
    '身高': 'height', '体重指数': 'bmi', 'BMI': 'bmi', '体重': 'weight',
    '收缩压': 'sbp', '舒张压': 'dbp',
    '电话': 'phone', '体检日期': 'date', '检查时间': 'time', '年龄': 'age',
}

def calc_age(idcard):
    import re
    m = re.match(r'\d{6}(\d{4})(\d{2})(\d{2})\d{3}[\dXx]', idcard)
    if not m: return None
    y, mo, d = int(m.group(1)), int(m.group(2)), int(m.group(3))
    from datetime import date
    today = date.today()
    age = today.year - y
    if (today.month, today.day) < (mo, d): age -= 1
    return str(age)

def detect_gender(idcard):
    """从身份证号判断性别：17位奇数为男，偶数为女"""
    import re
    m = re.match(r'\d{17}[\dXx]', idcard)
    if not m: return None
    digit = idcard[16]  # 第17位（0-based index 16）
    if digit in 'Xx': return None
    return '女' if int(digit) % 2 == 0 else '男'

def read_excel(path):
    import openpyxl
    wb = openpyxl.load_workbook(path)
    sheet = wb.active
    rows = list(sheet.iter_rows(values_only=True))
    if not rows: return []
    headers = [str(h).strip() if h else '' for h in rows[0]]
    col_index = {}
    for ci, h in enumerate(headers):
        best, best_len = None, 0
        for cn, en in COLUMN_MAP.items():
            if cn in h and len(cn) > best_len:
                best, best_len = en, len(cn)
        if best: col_index[best] = ci
    persons = []
    for row in rows[1:]:
        if all(v is None or str(v).strip() == '' for v in row): continue
        p = {}
        for en, ci in col_index.items():
            val = row[ci]
            if val is None: continue
            if isinstance(val, datetime): p[en] = val.strftime('%Y-%m-%d')
            elif isinstance(val, time): p[en] = val.strftime('%H:%M')
            else: p[en] = str(val).strip()
        if p: persons.append(p)
    return persons

def read_json(path):
    with open(path, 'r', encoding='utf-8') as f:
        return json.load(f)


def verify_docx(path, expected):
    """验证生成的docx文件中的字段是否与期望值一致。
    使用XML级搜索（可穿透文本框），与tihuan_tijian.py的分析逻辑一致。
    返回 (全部通过bool, 错误信息列表)"""
    import docx, re
    errors = []
    try:
        doc = docx.Document(path)
    except Exception as e:
        return False, [f"  无法打开文件: {e}"]

    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    # 合并段落文本（正文 + 表格 + 页眉页脚）
    texts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    texts.append(p.text)
    for section in doc.sections:
        for hf in [section.header, section.footer]:
            if hf:
                for p in hf.paragraphs:
                    texts.append(p.text)
    para_full = '\n'.join(texts)

    # 收集XML所有<w:t>元素文本（穿透文本框）
    all_t_texts = []
    # 正文body
    for t in doc._element.body.iter(f'{{{ns_w}}}t'):
        if t.text:
            all_t_texts.append(t.text)
    # 页眉页脚
    for section in doc.sections:
        for hf_obj in [section.header, section.first_page_header, section.even_page_header,
                        section.footer, section.first_page_footer, section.even_page_footer]:
            if hf_obj is None:
                continue
            for t in hf_obj._element.iter(f'{{{ns_w}}}t'):
                if t.text:
                    all_t_texts.append(t.text)
    xml_full = ''.join(all_t_texts)

    # 合并两种来源（段落文本保留空格，xml_full覆盖文本框）
    combined_text = para_full + '\n' + xml_full

    # 字段提取
    field_patterns = [
        ('number', r'体检[编代]码\s*[：:\t]?\s*(\d+)'),
        ('name',   r'姓\s*名\s*[：:\t]?\s*(\S+)'),
        ('age',    r'年\s*龄\s*[：:\t]?\s*(\d+)'),
        ('height', r'身\s*高\s*[：:\t]?\s*(\d+\.?\d*)'),
        ('weight', r'体\s*重\s*[：:\t]?\s*(\d+\.?\d*)'),
        ('bmi',    r'体重指数\s*[：:\t]?\s*(\d+\.?\d*)'),
        ('sbp',    r'收缩压\s*[：:\t]?\s*(\d+\.?\d*)'),
        ('dbp',    r'舒张压\s*[：:\t]?\s*(\d+\.?\d*)'),
        ('phone',  r'(1[3-9]\d{9})'),
    ]

    extracted = {}
    for key, pattern in field_patterns:
        m = re.search(pattern, combined_text)
        if m:
            extracted[key] = m.group(1)
        else:
            extracted[key] = ''

    # 身份证号：已知期望值，直接搜索期望文本是否存在于文档中
    exp_id = expected.get('idcard', '')
    if exp_id and exp_id in combined_text:
        extracted['idcard'] = exp_id
    elif exp_id:
        # 回退：搜索任何18位身份证号
        candidates = re.findall(r'(\d{17}[\dXx])', combined_text)
        if candidates:
            # 取第一个不以2026开头的（排除体检编码拼接）
            for c in candidates:
                if not c.startswith('2026'):
                    extracted['idcard'] = c
                    break
            else:
                extracted['idcard'] = candidates[0]
        else:
            extracted['idcard'] = ''
    else:
        extracted['idcard'] = ''

    # 体检日期：关键词和值可能在不同段落，搜索日期格式
    date_m = re.search(r'(\d{4}[-/]\d{1,2}[-/]\d{1,2})', combined_text)
    extracted['date'] = date_m.group(1) if date_m else ''

    # 对比
    check_fields = [
        ('name', '姓名'),
        ('idcard', '身份证号'),
        ('age', '年龄'),
        ('number', '体检编码'),
        ('height', '身高'),
        ('weight', '体重'),
        ('bmi', '体重指数'),
        ('sbp', '收缩压'),
        ('dbp', '舒张压'),
        ('phone', '电话'),
        ('date', '体检日期'),
    ]

    all_pass = True
    for key, label in check_fields:
        exp_val = expected.get(key, '')
        doc_val = extracted.get(key, '')
        if not exp_val:
            continue  # Excel中没有此字段则跳过
        exp_clean = str(exp_val).strip()
        doc_clean = str(doc_val).strip()
        # 数值字段去掉 .00 后缀
        if '.' in exp_clean and exp_clean.endswith('.00'):
            exp_clean = exp_clean.replace('.00', '')
        if '.' in doc_clean and doc_clean.endswith('.00'):
            doc_clean = doc_clean.replace('.00', '')
        if exp_clean == doc_clean:
            continue
        errors.append(f"  ❌ {label}: 文档值={doc_val}, 期望值={exp_val}")
        all_pass = False

    return all_pass, errors




if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("用法: python3 batch_tihuan.py persons.json|xlsx")
        sys.exit(1)
    path = sys.argv[1]
    if not os.path.isabs(path):
        path = os.path.join(WORK_DIR, path)
    ext = os.path.splitext(path)[1].lower()
    persons = read_excel(path) if ext in ('.xlsx', '.xls') else read_json(path) if ext == '.json' else []
    if not persons:
        print("❌ 没有读到人员信息"); sys.exit(1)

    for p in persons:
        if 'age' not in p and 'idcard' in p:
            age = calc_age(p['idcard'])
            if age: p['age'] = age
        for k in ('height', 'weight', 'bmi'):
            if k in p:
                try: p[k] = f"{float(p[k]):.2f}"
                except ValueError: pass

    print(f"📋 共 {len(persons)} 人")
    passed = []
    failed = []
    verified_errors = []

    for idx, p in enumerate(persons):
        name = p.get('name', f'person_{idx+1}')
        if not name or name.startswith('person_'):
            print(f"  ⚠️ [{idx+1}] 缺少姓名，跳过"); continue

        out = os.path.join(WORK_DIR, f"{name}.docx")
        shutil.copy(TEMPLATE, out)

        cmd = ["python3", SCRIPT, "--file", out]
        field_map = [
            ('name','--replace-name'),('idcard','--replace-idcard'),('age','--replace-age'),
            ('number','--replace-number'),('height','--replace-height'),('weight','--replace-weight'),
            ('bmi','--replace-bmi'),('sbp','--replace-sbp'),('dbp','--replace-dbp'),
            ('phone','--replace-phone'),('date','--replace-date'),('time','--replace-time'),
        ]
        for key, flag in field_map:
            if key in p and p[key]:
                cmd.extend([flag, str(p[key])])
        # 性别自动检测：身份证号17位奇男偶女
        if 'idcard' in p:
            gender = detect_gender(p['idcard'])
            if gender:
                cmd.extend(['--replace-gender', gender])

        r = subprocess.run(cmd, capture_output=True, text=True)
        if r.returncode != 0:
            print(f"  ❌ [{idx+1}/{len(persons)}] {name}.docx 替换失败")
            failed.append(name)
            continue

        # 自动验证
        v_pass, v_errors = verify_docx(out, p)
        if v_pass:
            print(f"  ✅ [{idx+1}/{len(persons)}] {name}.docx")
            passed.append(name)
        else:
            print(f"  ⚠️ [{idx+1}/{len(persons)}] {name}.docx 替换成功但验证发现问题:")
            for e in v_errors:
                print(e)
            verified_errors.append((name, v_errors))
            failed.append(name)

    # 清理备份
    for f in os.listdir(WORK_DIR):
        if 'backup_' in f:
            os.remove(os.path.join(WORK_DIR, f))

    # 汇总
    print(f"\n{'='*50}")
    print(f"📊 汇总")
    print(f"{'='*50}")
    if passed:
        print(f"✅ 成功: {len(passed)}/{len(persons)}")
        for n in passed:
            print(f"   ✅ {n}.docx")
    if verified_errors:
        print(f"\n❌ 验证失败: {len(verified_errors)}/{len(persons)}")
        for n, errs in verified_errors:
            print(f"   ❌ {n}.docx:")
            for e in errs:
                print(f"      {e}")
    if failed:
        print(f"\n❌ 替换失败: {len(failed)}/{len(persons)}")
        for n in failed:
            if n not in [x[0] for x in verified_errors]:
                print(f"   ❌ {n}.docx")

    if verified_errors or failed:
        print(f"\n⚠️ 存在错误！请检查上述问题后重新运行。")
        sys.exit(1)
    else:
        print(f"\n🎉 全部完成，所有文档已通过自动验证！")
