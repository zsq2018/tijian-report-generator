#!/usr/bin/env python3
"""全面验证每个生成的 docx 文件中的字段是否与 Excel 数据一致"""
import docx, re, openpyxl, os, sys
from datetime import date, datetime

WORK_DIR = os.path.dirname(os.path.abspath(__file__))

# 读取 Excel
wb = openpyxl.load_workbook(os.path.join(WORK_DIR, '人员信息模板.xlsx'))
ws = wb.active
rows = list(ws.iter_rows(values_only=True))
headers = list(rows[0])

COLUMN_MAP = {
    '姓名': 'name', '身份证号': 'idcard', '体检编码': 'number',
    '身高': 'height', '体重指数': 'bmi', 'BMI': 'bmi', '体重': 'weight',
    '收缩压': 'sbp', '舒张压': 'dbp',
    '电话': 'phone', '体检日期': 'date', '检查时间': 'time', '年龄': 'age',
}

def calc_age(idcard):
    m = re.match(r'\d{6}(\d{4})(\d{2})(\d{2})\d{3}[\dXx]', idcard)
    if not m: return None
    y, mo, d = int(m.group(1)), int(m.group(2)), int(m.group(3))
    today = date.today()
    age = today.year - y
    if (today.month, today.day) < (mo, d): age -= 1
    return str(age)

# 构建期望值列表
expects = []
for row in rows[1:]:
    if all(v is None or str(v).strip() == '' for v in row):
        continue
    person = {}
    for ci, h in enumerate(headers):
        best = None
        for cn, en in COLUMN_MAP.items():
            if cn == h:
                best = en
                break
        if best:
            val = row[ci]
            if isinstance(val, datetime):
                person[best] = val.strftime('%Y-%m-%d')
            elif isinstance(val, (int, float)):
                person[best] = str(val)
            elif val:
                person[best] = str(val).strip()
    if person.get('name'):
        for k in ('height', 'weight', 'bmi'):
            if k in person:
                try:
                    person[k] = f'{float(person[k]):.2f}'
                except:
                    pass
        if 'age' not in person and 'idcard' in person:
            age = calc_age(person['idcard'])
            if age:
                person['age'] = age
        expects.append(person)

print(f'📋 Excel 数据（共 {len(expects)} 人）：')
for p in expects:
    print(f'  {p["name"]}: 编码={p.get("number","")}, 身份证={p.get("idcard","")}, '
          f'年龄={p.get("age","")}, 身高={p.get("height","")}, 体重={p.get("weight","")}, '
          f'BMI={p.get("bmi","")}, 血压={p.get("sbp","")}/{p.get("dbp","")}, '
          f'电话={p.get("phone","")}, 日期={p.get("date","")}, 时间={p.get("time","")}')

ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
pass_all = 0
fail_all = 0

for person in expects:
    name = person['name']
    docx_path = os.path.join(WORK_DIR, f'{name}.docx')
    print(f'\n{"="*60}')
    print(f'📋 [{name}]')
    print(f'{"="*60}')

    if not os.path.exists(docx_path):
        print(f'  ❌ 文件不存在！')
        fail_all += 1
        continue

    try:
        doc = docx.Document(docx_path)
    except Exception as e:
        print(f'  ❌ 无法打开: {e}')
        fail_all += 1
        continue

    # 段落文本
    para_texts = [para.text for para in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para_texts.append(para.text)
    for section in doc.sections:
        for hf in [section.header, section.footer]:
            if hf:
                for para in hf.paragraphs:
                    para_texts.append(para.text)
    para_full = '\n'.join(para_texts)

    # XML 级所有 <w:t> 文本（穿透文本框）
    all_t_texts = []
    for t in doc._element.body.iter(f'{{{ns_w}}}t'):
        if t.text:
            all_t_texts.append(t.text)
    for section in doc.sections:
        for hf_obj in [section.header, section.first_page_header, section.even_page_header,
                       section.footer, section.first_page_footer, section.even_page_footer]:
            if hf_obj is None:
                continue
            for t in hf_obj._element.iter(f'{{{ns_w}}}t'):
                if t.text:
                    all_t_texts.append(t.text)
    xml_full = ''.join(all_t_texts)
    combined = para_full + '\n' + xml_full

    # 提取字段
    extracted = {}
    field_patterns = [
        ('number', r'体检[编代]码\s*[：:\t]?\s*(\d+)'),
        ('name',   r'姓\s*名\s*[：:\t]?\s*(\S+)'),
        ('age',    r'年\s*龄\s*[：:\t]?\s*(\d+)'),
        ('height', r'身\s*高\s*[：:\t]?\s*(\d+\.?\d*)'),
        ('weight', r'体\s*重\s*[：:\t]?\s*(\d+\.?\d*)'),
        ('bmi',    r'体重指数\s*[：:\t]?\s*(\d+\.?\d*)'),
        ('sbp',    r'收缩压\s*[：:\t]?\s*(\d+\.?\d*)'),
        ('dbp',    r'舒张压\s*[：:\t]?\s*(\d+\.?\d*)'),
        ('phone',  r'(1[3-9]\d{9})'),
    ]
    for key, pat in field_patterns:
        m = re.search(pat, combined)
        extracted[key] = m.group(1) if m else ''

    # 身份证号
    exp_id = person.get('idcard', '')
    if exp_id and exp_id in combined:
        extracted['idcard'] = exp_id
    elif exp_id:
        cands = re.findall(r'(\d{17}[\dXx])', combined)
        extracted['idcard'] = ''
        for c in cands:
            if not c.startswith('2026'):
                extracted['idcard'] = c
                break
        if not extracted['idcard'] and cands:
            extracted['idcard'] = cands[0]
    else:
        extracted['idcard'] = ''

    dm = re.search(r'(\d{4}[-/]\d{1,2}[-/]\d{1,2})', combined)
    extracted['date'] = dm.group(1) if dm else ''
    tm = re.search(r'(\d{1,2}:\d{2})', combined)
    extracted['time'] = tm.group(1) if tm else ''

    check_fields = [
        ('name', '姓名'), ('idcard', '身份证号'), ('age', '年龄'),
        ('number', '体检编码'), ('height', '身高'), ('weight', '体重'),
        ('bmi', '体重指数'), ('sbp', '收缩压'), ('dbp', '舒张压'),
        ('phone', '电话'), ('date', '体检日期'), ('time', '检查时间'),
    ]

    all_ok = True
    for key, label in check_fields:
        exp_val = person.get(key, '')
        doc_val = extracted.get(key, '')
        if not exp_val:
            continue
        exp_clean = str(exp_val).strip()
        doc_clean = str(doc_val).strip()
        if '.' in exp_clean and exp_clean.endswith('.00'):
            exp_clean = exp_clean.replace('.00', '')
        if '.' in doc_clean and doc_clean.endswith('.00'):
            doc_clean = doc_clean.replace('.00', '')
        if exp_clean == doc_clean:
            print(f'  ✅ {label}: {doc_val}')
        else:
            print(f'  ❌ {label}: 文档值="{doc_val}", 期望值="{exp_val}"')
            all_ok = False

    if all_ok:
        print(f'  ✅ 全部字段匹配！')
        pass_all += 1
    else:
        fail_all += 1

print(f'\n{"="*60}')
print(f'📊 最终汇总')
print(f'{"="*60}')
print(f'  通过: {pass_all}/{len(expects)}')
print(f'  失败: {fail_all}/{len(expects)}')
if fail_all == 0:
    print(f'  🎉 全部文档替换正确！')
else:
    print(f'  ❌ 存在错误，请检查！')

sys.exit(0 if fail_all == 0 else 1)
