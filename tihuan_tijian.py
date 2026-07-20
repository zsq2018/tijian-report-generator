#!/usr/bin/env python3
"""
体检报告全局替换工具
=========================
逻辑：识别文档中关键词后面的值 → 全局替换所有相同的值
      （正文 + 表格 + 页眉页脚含文本框）

用法：
  python3 tihuan_tijian.py                                    # 分析模式
  python3 tihuan_tijian.py --dry-run                          # 预览
  python3 tihuan_tijian.py --replace-number 新值              # 替换体检编码
  python3 tihuan_tijian.py --replace-name 新值                 # 替换姓名
  python3 tihuan_tijian.py --replace-age 新值                  # 替换年龄
  python3 tihuan_tijian.py --replace-phone 新值                # 替换电话
  python3 tihuan_tijian.py --replace-date 新值                 # 替换体检日期
  python3 tihuan_tijian.py --replace-time 新值                 # 替换检查时间
  python3 tihuan_tijian.py --replace-height 新值               # 替换身高
  python3 tihuan_tijian.py --replace-weight 新值               # 替换体重
  python3 tihuan_tijian.py --replace-bmi 新值                  # 替换体重指数
  python3 tihuan_tijian.py --replace-sbp 新值                  # 替换收缩压
  python3 tihuan_tijian.py --replace-dbp 新值                  # 替换舒张压
  （以上参数可任意组合）
"""

import docx
import re
import sys
import shutil
import os
from datetime import datetime
from lxml import etree


# ============================================================
#  替换项配置（关键词正则 → 参数名 → 标签）
# ============================================================

class Item:
    def __init__(self, key, label, pattern, finder=None):
        self.key = key          # 参数名，如 'number', 'name'
        self.label = label      # 中文标签，如 '体检编码'
        self.pattern = pattern  # 用于在段落/表格文本中匹配关键词+值的正则
        self.finder = finder    # 可选的自定义查找函数，优先级高于 pattern

    def find_in_text(self, text):
        """在文本中匹配，返回 (值, 起始位置, 结束位置) 或 None"""
        m = re.search(self.pattern, text)
        if m:
            return m.group(1), m.start(1), m.end(1)
        return None

    def find_in_paragraphs(self, paragraphs, loc_prefix=""):
        """在段落列表中查找关键词"""
        for pi, para in enumerate(paragraphs):
            result = self.find_in_text(para.text)
            if result:
                val, _, _ = result
                return val, f"{loc_prefix}段落{pi}", para.text.strip()[:80]
        return None, None, None


# --- 简单关键词匹配的项 ---
SIMPLE_ITEMS = [
    Item('number', '体检编码', r'体检[编代]码\s*[：:\t]?\s*(\d+)'),
    Item('name',   '姓名',    r'姓\s*名\s*[：:\t]?\s*(\S+)'),
    Item('age',    '年龄',    r'年\s*龄\s*[：:\t]?\s*(\d+)'),
    Item('height', '身高',    r'身\s*高\s*[：:\t]?\s*(\d+\.?\d*)'),
    Item('weight', '体重',    r'体\s*重\s*[：:\t]?\s*(\d+\.?\d*)'),
    Item('bmi',    '体重指数', r'体重指数\s*[：:\t]?\s*(\d+\.?\d*)'),
    Item('sbp',    '收缩压',  r'收缩压\s*[：:\t]?\s*(\d+\.?\d*)'),
    Item('dbp',    '舒张压',  r'舒张压\s*[：:\t]?\s*(\d+\.?\d*)'),
]


# --- 需要跨段落查找的项 ---
def _find_phone(doc):
    """联系电话：关键词和号码可能在不同段落"""
    for i, para in enumerate(doc.paragraphs):
        m = re.search(r'联系电话\s*[：:\t]?\s*(1[3-9]\d{9})', para.text)
        if m:
            return m.group(1), f"段落{i}", para.text.strip()[:80]
    for i, para in enumerate(doc.paragraphs):
        if '联系电话' in para.text:
            for j in range(i, min(i + 4, len(doc.paragraphs))):
                m = re.search(r'(1[3-9]\d{9})', doc.paragraphs[j].text)
                if m:
                    return m.group(1), f"段落{j}（关键词在段落{i}）", doc.paragraphs[j].text.strip()[:80]
    return _search_regex(doc, r'联系电话\s*[：:\t]?\s*(1[3-9]\d{9})')


def _find_date(doc):
    """体检日期：关键词和日期可能在不同段落"""
    for i, para in enumerate(doc.paragraphs):
        m = re.search(r'体检日期\s*[：:\t]?\s*(\d{4}[-/]\d{1,2}[-/]\d{1,2})', para.text)
        if m:
            return m.group(1), f"段落{i}", para.text.strip()[:80]
    for i, para in enumerate(doc.paragraphs):
        if '体检日期' in para.text:
            for j in range(i, min(i + 4, len(doc.paragraphs))):
                m = re.search(r'(\d{4}[-/]\d{1,2}[-/]\d{1,2})', doc.paragraphs[j].text)
                if m:
                    return m.group(1), f"段落{j}（关键词在段落{i}）", doc.paragraphs[j].text.strip()[:80]
    return _search_regex(doc, r'体检日期\s*[：:\t]?\s*(\d{4}[-/]\d{1,2}[-/]\d{1,2})')


def _find_time(doc):
    """检查时间：关键词和时间可能在不同位置"""
    pattern = r'(?:检查日期|报告日期)\s*[：:\t]?\s*\d{4}[-/]\d{1,2}[-/]\d{1,2}\s+(\d{1,2}:\d{2})'
    for i, para in enumerate(doc.paragraphs):
        m = re.search(pattern, para.text)
        if m:
            return m.group(1), f"段落{i}", para.text.strip()[:80]
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                m = re.search(pattern, cell.text)
                if m:
                    return m.group(1), f"表格{ti}行{ri}列{ci}", cell.text.strip()[:80]
    for label in ['检查日期', '报告日期']:
        for i, para in enumerate(doc.paragraphs):
            if label in para.text:
                for j in range(i, min(i + 4, len(doc.paragraphs))):
                    m = re.search(r'(\d{1,2}:\d{2})', doc.paragraphs[j].text)
                    if m:
                        return m.group(1), f"段落{j}（关键词在段落{i}）", doc.paragraphs[j].text.strip()[:80]
    for label in ['检查日期', '报告日期']:
        for ti, table in enumerate(doc.tables):
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    if label in cell.text:
                        m = re.search(r'(\d{1,2}:\d{2})', cell.text)
                        if m:
                            return m.group(1), f"表格{ti}行{ri}列{ci}", cell.text.strip()[:80]
    return None, None, None


def _find_idcard(doc):
    """查找18位身份证号 —— 关键词和值可能在不同段落/表格中"""
    import re
    # 正文段落中搜索
    for i, para in enumerate(doc.paragraphs):
        if '身份证号' not in para.text:
            continue
        # 在当前段落关键词后面找
        kw_m = re.search(r'身份证号', para.text)
        tail = para.text[kw_m.end():]
        id_m = re.search(r'(\d{17}[\dXx])', tail)
        if id_m:
            return id_m.group(1), f"段落{i}", para.text.strip()[:80]
        # 跨段落：搜索后续段落
        for j in range(i + 1, min(i + 6, len(doc.paragraphs))):
            next_text = doc.paragraphs[j].text
            id_m2 = re.search(r'(\d{17}[\dXx])', next_text)
            if id_m2:
                return id_m2.group(1), f"段落{j}（关键词在段落{i}）", next_text.strip()[:80]
    # 表格中搜索
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                if '身份证号' not in cell.text:
                    continue
                kw_m = re.search(r'身份证号', cell.text)
                tail = cell.text[kw_m.end():]
                id_m = re.search(r'(\d{17}[\dXx])', tail)
                if id_m:
                    return id_m.group(1), f"表格{ti}行{ri}列{ci}", cell.text.strip()[:80]
                # 跨单元格/跨行搜索
                cell_paras = cell.paragraphs
                for pidx, cp in enumerate(cell_paras):
                    if '身份证号' in cp.text:
                        # 搜同一格后续段落
                        for k in range(pidx + 1, len(cell_paras)):
                            id_m2 = re.search(r'(\d{17}[\dXx])', cell_paras[k].text)
                            if id_m2:
                                return id_m2.group(1), f"表格{ti}行{ri}列{ci}", cell_paras[k].text.strip()[:80]
                        # 搜同行后续列
                        for cj in range(ci + 1, len(row.cells)):
                            for cp2 in row.cells[cj].paragraphs:
                                id_m2 = re.search(r'(\d{17}[\dXx])', cp2.text)
                                if id_m2:
                                    return id_m2.group(1), f"表格{ti}行{ri}列{cj}", cp2.text.strip()[:80]
    # XML 回退：搜索文本框（python-docx 不直接暴露文本框内文本）
    import re
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    # 1. 搜索所有文本框中的 w:t 元素（标签和值可能在不同文本框，需跨文本框搜索）
    all_telems = []
    for txbx in doc._element.body.iter(f'{{{ns_w}}}txbxContent'):
        for t in txbx.iter(f'{{{ns_w}}}t'):
            all_telems.append(t)
    for section in doc.sections:
        for hf_obj in [section.header, section.first_page_header, section.even_page_header,
                        section.footer, section.first_page_footer, section.even_page_footer]:
            if hf_obj is None:
                continue
            for txbx in hf_obj._element.iter(f'{{{ns_w}}}txbxContent'):
                for t in txbx.iter(f'{{{ns_w}}}t'):
                    all_telems.append(t)

    for idx, t in enumerate(all_telems):
        if '身份证号' not in (t.text or ''):
            continue
        # 先查当前元素关键词后面
        kw_pos = (t.text or '').find('身份证号')
        tail = (t.text or '')[kw_pos + 4:]
        id_m = re.search(r'(\d{17}[\dXx])', tail)
        if id_m:
            return id_m.group(1), '正文（文本框）', ''
        # 逐个检查后续元素（不拼接，先找完整独立的值）
        for j in range(idx + 1, min(idx + 20, len(all_telems))):
            elem_text = (all_telems[j].text or '').strip()
            id_m = re.match(r'^(\d{17}[\dXx])$', elem_text)
            if id_m:
                return id_m.group(1), '正文（文本框）', ''
        # 回退：拼接后续纯数字元素
        combined = ''
        for j in range(idx + 1, min(idx + 20, len(all_telems))):
            elem_text = (all_telems[j].text or '')
            if re.match(r'^[\dXx]+$', elem_text):
                combined += elem_text
                id_m = re.search(r'(\d{17}[\dXx])', combined)
                if id_m:
                    return id_m.group(1), '正文（文本框）', ''
            elif combined:
                break
        break

    # 3. 最后回退：限制范围的全 body XML 搜索（只在关键词后 2000 字符内查找，减少跨区域误匹配）
    body_xml = etree.tostring(doc._element.body, encoding='unicode')
    kw_pos = body_xml.find('身份证号')
    if kw_pos >= 0:
        tail_xml = body_xml[kw_pos:kw_pos + 2000]
        tail_plain = re.sub(r'<[^>]+>', '', tail_xml)
        id_m = re.search(r'(\d{17}[\dXx])', tail_plain)
        if id_m:
            return id_m.group(1), '正文（XML回退）', tail_plain[:80]
    return None, None, None


def _find_gender(doc):
    """查找性别：查找关键词后的 男/女"""
    keywords = ['性\u0009别', '性 别', '性别', '姓别']
    # 段落中查找
    for i, para in enumerate(doc.paragraphs):
        for kw in keywords:
            if kw not in para.text:
                continue
            kw_m = re.search(re.escape(kw), para.text)
            tail = para.text[kw_m.end():]
            g_m = re.search(r'\s*([男女])', tail)
            if g_m:
                return g_m.group(1), f"段落{i}", para.text.strip()[:80]
    # 表格中查找
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for kw in keywords:
                    if kw not in cell.text:
                        continue
                    kw_m = re.search(re.escape(kw), cell.text)
                    tail = cell.text[kw_m.end():]
                    g_m = re.search(r'\s*([男女])', tail)
                    if g_m:
                        return g_m.group(1), f"表格{ti}行{ri}列{ci}", cell.text.strip()[:80]
    # 文本框查找
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    all_telems = []
    for txbx in doc._element.body.iter(f'{{{ns_w}}}txbxContent'):
        for t in txbx.iter(f'{{{ns_w}}}t'):
            all_telems.append(t)
    for section in doc.sections:
        for hf_obj in [section.header, section.first_page_header, section.even_page_header,
                        section.footer, section.first_page_footer, section.even_page_footer]:
            if hf_obj is None:
                continue
            for txbx in hf_obj._element.iter(f'{{{ns_w}}}txbxContent'):
                for t in txbx.iter(f'{{{ns_w}}}t'):
                    all_telems.append(t)
    for idx, t in enumerate(all_telems):
        for kw in keywords:
            if kw not in (t.text or ''):
                continue
            kw_pos = (t.text or '').find(kw)
            tail = (t.text or '')[kw_pos + len(kw):]
            g_m = re.search(r'\s*([男女])', tail)
            if g_m:
                return g_m.group(1), '文本框', ''
            # 后续元素
            for j in range(idx + 1, min(idx + 5, len(all_telems))):
                g_m2 = re.search(r'([男女])', (all_telems[j].text or ''))
                if g_m2:
                    return g_m2.group(1), '文本框', ''
            break
    # XML回退：在关键词后2000字符内查找
    body_xml = etree.tostring(doc._element.body, encoding='unicode')
    for kw in keywords:
        kw_pos = body_xml.find(kw)
        if kw_pos >= 0:
            tail_xml = body_xml[kw_pos:kw_pos + 500]
            tail_plain = re.sub(r'<[^>]+>', '', tail_xml)
            g_m = re.search(r'([男女])', tail_plain)
            if g_m:
                return g_m.group(1), 'XML回退', tail_plain[:80]
    return None, None, None


def replace_gender_in_doc(doc, old_gen, new_gen):
    """替换文档中关键词后的性别（男/女）"""
    keywords = ['性\u0009别', '性 别', '性别', '姓别', '性别：']
    count = 0

    # 1. 替换段落中的性别
    for para in doc.paragraphs:
        for kw in keywords:
            if kw not in para.text:
                continue
            kw_m = re.search(re.escape(kw), para.text)
            tail = para.text[kw_m.end():]
            g_m = re.search(r'\s*' + re.escape(old_gen), tail)
            if g_m:
                abs_start = kw_m.end() + g_m.start() + (g_m.end() - g_m.start() - len(old_gen))
                result = _replace_at_runs_position(para, abs_start, abs_start + len(old_gen), new_gen)
                if result:
                    count += 1
                break

    # 2. 替换表格中的性别
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for kw in keywords:
                        if kw not in para.text:
                            continue
                        kw_m = re.search(re.escape(kw), para.text)
                        tail = para.text[kw_m.end():]
                        g_m = re.search(r'\s*' + re.escape(old_gen), tail)
                        if g_m:
                            abs_start = kw_m.end() + g_m.start() + (g_m.end() - g_m.start() - len(old_gen))
                            result = _replace_at_runs_position(para, abs_start, abs_start + len(old_gen), new_gen)
                            if result:
                                count += 1
                            break

    # 3. 替换文本框中的性别（关键词后的性别值）
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    ns_t = f'{{{ns_w}}}t'
    all_telems = []
    for txbx in doc._element.body.iter(f'{{{ns_w}}}txbxContent'):
        for t in txbx.iter(ns_t):
            all_telems.append(t)
    for section in doc.sections:
        for hf_obj in [section.header, section.first_page_header, section.even_page_header,
                        section.footer, section.first_page_footer, section.even_page_footer]:
            if hf_obj is None:
                continue
            for txbx in hf_obj._element.iter(f'{{{ns_w}}}txbxContent'):
                for t in txbx.iter(ns_t):
                    all_telems.append(t)

    search_start = 0
    while True:
        joined = ''.join(t.text or '' for t in all_telems)
        # 找任意关键词
        best_kw, best_pos = None, len(joined)
        for kw in keywords:
            pos = joined.find(kw, search_start)
            if 0 <= pos < best_pos:
                best_kw, best_pos = kw, pos
        if best_kw is None:
            break
        kw_end = best_pos + len(best_kw)
        tail = joined[kw_end:]
        g_m = re.search(r'\s*' + re.escape(old_gen), tail)
        if not g_m:
            search_start = kw_end
            continue
        match_text = g_m.group(0)  # e.g. " 男"
        gen_pos_in_match = match_text.find(old_gen)
        abs_start = kw_end + g_m.start() + gen_pos_in_match
        abs_end = abs_start + len(old_gen)

        # 找到对应元素并替换
        length = 0
        start_t = 0
        for i, t in enumerate(all_telems):
            tl = len(t.text or '')
            if length + tl > abs_start:
                start_t = i
                break
            length += tl
        else:
            search_start = kw_end + 1
            continue
        off = abs_start - sum(len(all_telems[j].text or '') for j in range(start_t))
        end_t = start_t
        for i in range(start_t, len(all_telems)):
            if sum(len(all_telems[j].text or '') for j in range(i + 1)) >= abs_end:
                end_t = i
                break
        if start_t == end_t:
            all_telems[start_t].text = (all_telems[start_t].text or '')[:off] + new_gen + (all_telems[start_t].text or '')[off + len(old_gen):]
        else:
            all_telems[start_t].text = (all_telems[start_t].text or '')[:off] + new_gen
            for i in range(start_t + 1, end_t):
                all_telems[i].text = ''
            remain = sum(len(all_telems[j].text or '') for j in range(end_t)) - abs_end
            all_telems[end_t].text = (all_telems[end_t].text or '')[-remain:] if remain > 0 else ''
        count += 1
        search_start = abs_end
        # 重新加载 all_elems 因为修改了 text
        # 实际上不需要，因为 same elements

    return count


CUSTOM_ITEMS = [
    Item('idcard', '身份证号', None, finder=_find_idcard),
    Item('phone', '联系电话', None, finder=_find_phone),
    Item('date',  '体检日期', None, finder=_find_date),
    Item('time',  '检查时间', None, finder=_find_time),
    Item('gender','性别',    None, finder=_find_gender),
]

ALL_ITEMS = SIMPLE_ITEMS + CUSTOM_ITEMS


def _search_regex(doc, pattern):
    """在段落/表格中用正则搜索"""
    for i, para in enumerate(doc.paragraphs):
        m = re.search(pattern, para.text)
        if m:
            return m.group(1), f"段落{i}", para.text.strip()[:80]
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                m = re.search(pattern, cell.text)
                if m:
                    return m.group(1), f"表格{ti}行{ri}列{ci}", cell.text.strip()[:80]
    return None, None, None


# ============================================================
#  查找函数（统一入口）
# ============================================================


def calc_age_from_idcard(idcard, ref_date=None):
    """从18位身份证号计算年龄"""
    import re
    m = re.match(r'\d{6}(\d{4})(\d{2})(\d{2})\d{3}[\dXx]', idcard)
    if not m:
        return None
    birth_y, birth_m, birth_d = int(m.group(1)), int(m.group(2)), int(m.group(3))
    if ref_date is None:
        from datetime import date
        ref_date = date.today()
    age = ref_date.year - birth_y
    if (ref_date.month, ref_date.day) < (birth_m, birth_d):
        age -= 1
    return str(age)

def find_target(doc, item):
    """根据 Item 配置查找目标值（含 XML 回退）"""
    if item.finder:
        return item.finder(doc)
    # 先在段落中找
    for i, para in enumerate(doc.paragraphs):
        result = item.find_in_text(para.text)
        if result:
            return result[0], f"段落{i}", para.text.strip()[:80]
    # 再在表格中找
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                result = item.find_in_text(cell.text)
                if result:
                    return result[0], f"表格{ti}行{ri}列{ci}", cell.text.strip()[:80]
    # XML 回退：在 raw body XML 中用宽泛模式搜索
    import re
    body_xml = etree.tostring(doc._element.body, encoding='unicode')
    # 从 pattern 中提取关键词和目标值正则
    # pattern 形如 r'关键词...(\d+)' -> 拆为关键词和目标模式
    parts = item.pattern.rsplit('(', 1)
    if len(parts) == 2:
        kw_pattern = parts[0]  # 关键词部分（含正则）
        val_pattern = '(' + parts[1]  # 值捕获组
        # 放宽关键词模式：将 \s* 替换为 [\s<][^>]*> 来跨越XML标签
        relaxed_kw = kw_pattern.replace(r'\s*', r'[\s<][^>]*>')
        relaxed_kw = relaxed_kw.replace(r'	', r'[	<][^>]*>')
        relaxed_kw = relaxed_kw.replace(r'[：:	]?', r'[：:	<][^>]*>')
        relaxed_pattern = relaxed_kw + r'[\s\S]*?' + val_pattern
        m = re.search(relaxed_pattern, body_xml)
        if m:
            return m.group(1), '正文（XML回退）', ''
    return None, None, None


# ============================================================
#  匹配收集
# ============================================================

def count_matches_in_paragraphs(paragraphs, target, loc_prefix=""):
    results = []
    for pi, para in enumerate(paragraphs):
        if target not in para.text:
            continue
        runs = para.runs
        rt_list = [r.text for r in runs]
        cmap = []
        for ri, rt in enumerate(rt_list):
            cmap.extend([ri] * len(rt))
        pos = 0
        while True:
            idx = para.text.find(target, pos)
            if idx == -1:
                break
            end_idx = idx + len(target)
            sr = cmap[idx] if idx < len(cmap) else 0
            er = cmap[end_idx - 1] if end_idx - 1 < len(cmap) else len(runs) - 1
            run_info = f"run{sr}" if sr == er else f"run{sr}→run{er}（跨run）"
            results.append({
                'location': f"{loc_prefix}段落{pi}, {run_info}",
                'context': para.text.strip()[:80],
                'run_range': (sr, er),
            })
            pos = end_idx
    return results


def _count_in_xml(element, target):
    return etree.tostring(element, encoding='unicode').count(target)


def collect_all_matches(doc, target):
    results = []
    results += count_matches_in_paragraphs(doc.paragraphs, target)
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                results += count_matches_in_paragraphs(
                    cell.paragraphs, target,
                    loc_prefix=f"表格{ti}行{ri}列{ci}内"
                )
    seen_hf = set()
    for si, section in enumerate(doc.sections):
        for hf_name, hf_obj in [
            ('header', section.header), ('first_page_header', section.first_page_header),
            ('even_page_header', section.even_page_header),
            ('footer', section.footer), ('first_page_footer', section.first_page_footer),
            ('even_page_footer', section.even_page_footer),
        ]:
            if hf_obj is None:
                continue
            elem_id = id(hf_obj._element)
            if elem_id in seen_hf:
                continue
            seen_hf.add(elem_id)
            results += count_matches_in_paragraphs(hf_obj.paragraphs, target, loc_prefix=f"节{si}{hf_name}内")
            xml_count = _count_in_xml(hf_obj._element, target)
            para_count = sum(1 for p in hf_obj.paragraphs if target in p.text)
            extra = xml_count - para_count
            for _ in range(extra):
                results.append({'location': f"节{si}{hf_name}内（文本框）", 'context': '(文本框内)', 'run_range': (0, 0)})
    return results


# ============================================================
#  替换函数
# ============================================================

def replace_in_paragraphs(paragraphs, old_text, new_text):
    count = 0
    for para in paragraphs:
        if old_text not in para.text:
            continue
        runs = para.runs
        rtl = [r.text for r in runs]
        full = ''.join(rtl)
        pos = 0
        reps = []
        while True:
            idx = full.find(old_text, pos)
            if idx == -1:
                break
            reps.append((idx, idx + len(old_text)))
            pos = idx + len(old_text)
            count += 1
        for sc, ec in reversed(reps):
            rtl = [r.text for r in runs]
            cmap = []
            for ri, rt in enumerate(rtl):
                cmap.extend([ri] * len(rt))
            offs = [0]
            for rt in rtl:
                offs.append(offs[-1] + len(rt))
            sr = cmap[sc]
            er = cmap[ec - 1]
            os_ = sc - offs[sr]
            oe_ = ec - offs[er]
            if sr == er:
                runs[sr].text = runs[sr].text[:os_] + new_text + runs[sr].text[oe_:]
            else:
                runs[sr].text = runs[sr].text[:os_] + new_text
                for ri in range(sr + 1, er):
                    runs[ri].text = ''
                runs[er].text = runs[er].text[oe_:]
    return count


def _replace_in_xml_element(element, old_text, new_text):
    count = 0
    for t_elem in element.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
        if t_elem.text and old_text in t_elem.text:
            occurrences = t_elem.text.count(old_text)
            t_elem.text = t_elem.text.replace(old_text, new_text)
            count += occurrences
    return count


def replace_all(doc, old_text, new_text):
    total = replace_in_paragraphs(doc.paragraphs, old_text, new_text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                total += replace_in_paragraphs(cell.paragraphs, old_text, new_text)
    for section in doc.sections:
        for hf in [section.header, section.first_page_header, section.even_page_header,
                    section.footer, section.first_page_footer, section.even_page_footer]:
            if hf is None:
                continue
            total += replace_in_paragraphs(hf.paragraphs, old_text, new_text)
            total += _replace_in_xml_element(hf._element, old_text, new_text)
    # 对整篇文档 body 做 XML 级替换（覆盖段落API可能遗漏的内容）
    total += _replace_in_xml_element(doc._element.body, old_text, new_text)
    return total


# ============================================================
#  精准替换函数（仅替换关键词后面的值，不全局替换）
# ============================================================

def _replace_at_runs_position(para, start, end, new_text):
    """在段落的 runs 中替换指定位置范围的文本。返回替换次数。"""
    runs = para.runs
    rtl = [r.text for r in runs]
    cmap = []
    for ri, rt in enumerate(rtl):
        cmap.extend([ri] * len(rt))
    if start >= len(cmap):
        return 0
    sr = cmap[start]
    er = cmap[end - 1] if end - 1 < len(cmap) else len(runs) - 1
    offs = [0]
    for rt in rtl:
        offs.append(offs[-1] + len(rt))
    os_ = start - offs[sr]
    oe_ = end - offs[er]
    if sr == er:
        runs[sr].text = runs[sr].text[:os_] + new_text + runs[sr].text[oe_:]
    else:
        runs[sr].text = runs[sr].text[:os_] + new_text
        for ri in range(sr + 1, er):
            runs[ri].text = ''
        runs[er].text = runs[er].text[oe_:]
    return 1


def replace_keyword_only(doc, keyword, old_val, new_val):
    """只替换关键词后面的值，不进行全局替换。支持跨段落搜索。返回替换次数。"""
    import re
    escaped_kw = re.escape(keyword)
    escaped_old = re.escape(old_val)

    # 搜索正文段落（含跨段落）
    for i, para in enumerate(doc.paragraphs):
        kw_m = re.search(escaped_kw, para.text)
        if not kw_m:
            continue
        # 当前段落关键词后面
        tail = para.text[kw_m.end():]
        val_m = re.search(escaped_old, tail)
        if val_m:
            val_start = kw_m.end() + val_m.start()
            val_end = val_start + len(old_val)
            return _replace_at_runs_position(para, val_start, val_end, new_val)
        # 跨段落：搜索后续段落
        for j in range(i + 1, min(i + 6, len(doc.paragraphs))):
            next_para = doc.paragraphs[j]
            val_m2 = re.search(escaped_old, next_para.text)
            if val_m2:
                return _replace_at_runs_position(next_para, val_m2.start(), val_m2.end(), new_val)

    # 搜索表格（含跨单元格）
    for table in doc.tables:
        for row in table.rows:
            for ci, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    kw_m = re.search(escaped_kw, para.text)
                    if not kw_m:
                        continue
                    tail = para.text[kw_m.end():]
                    val_m = re.search(escaped_old, tail)
                    if val_m:
                        val_start = kw_m.end() + val_m.start()
                        val_end = val_start + len(old_val)
                        return _replace_at_runs_position(para, val_start, val_end, new_val)
                    # 同一格后续段落
                    cell_paras = cell.paragraphs
                    for pidx, cp in enumerate(cell_paras):
                        if keyword not in cp.text:
                            continue
                        for k in range(pidx + 1, len(cell_paras)):
                            val_m2 = re.search(escaped_old, cell_paras[k].text)
                            if val_m2:
                                return _replace_at_runs_position(cell_paras[k], val_m2.start(), val_m2.end(), new_val)
                    # 同行后续列
                    for cj in range(ci + 1, len(row.cells)):
                        for cp2 in row.cells[cj].paragraphs:
                            val_m2 = re.search(escaped_old, cp2.text)
                            if val_m2:
                                return _replace_at_runs_position(cp2, val_m2.start(), val_m2.end(), new_val)

    # XML 回退：跨所有文本框搜索关键词并替换（标签和值可能在不同文本框）
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    ns_t = f'{{{ns_w}}}t'

    # 收集所有文本框的 w:t 元素（含正文和页眉页脚）
    all_telems = []
    for txbx in doc._element.body.iter(f'{{{ns_w}}}txbxContent'):
        for t in txbx.iter(ns_t):
            all_telems.append(t)
    for section in doc.sections:
        for hf in [section.header, section.first_page_header, section.even_page_header,
                    section.footer, section.first_page_footer, section.even_page_footer]:
            if hf is None:
                continue
            for txbx in hf._element.iter(f'{{{ns_w}}}txbxContent'):
                for t in txbx.iter(ns_t):
                    all_telems.append(t)

    # 循环替换：可能有多个文本框包含相同关键词+值，需全部替换
    count = 0
    search_start = 0
    while True:
        joined = ''.join(t.text or '' for t in all_telems)
        kw_idx = joined.find(keyword, search_start)
        if kw_idx < 0:
            break
        tail = joined[kw_idx + len(keyword):]
        val_idx = tail.find(old_val)
        if val_idx < 0:
            search_start = kw_idx + len(keyword)  # 跳过此关键词，继续找下一个
            continue
        abs_start = kw_idx + len(keyword) + val_idx
        abs_end = abs_start + len(old_val)

        # 跨元素替换
        length = 0
        start_t = 0
        for i, t in enumerate(all_telems):
            tl = len(t.text or '')
            if length + tl > abs_start:
                start_t = i
                break
            length += tl
        else:
            search_start = kw_idx + len(keyword)
            continue
        off = abs_start - sum(len(all_telems[j].text or '') for j in range(start_t))
        end_t = start_t
        for i in range(start_t, len(all_telems)):
            if sum(len(all_telems[j].text or '') for j in range(i + 1)) >= abs_end:
                end_t = i
                break
        if start_t == end_t:
            all_telems[start_t].text = (all_telems[start_t].text or '')[:off] + new_val + (all_telems[start_t].text or '')[off + len(old_val):]
        else:
            all_telems[start_t].text = (all_telems[start_t].text or '')[:off] + new_val
            for i in range(start_t + 1, end_t):
                all_telems[i].text = ''
            remain = sum(len(all_telems[j].text or '') for j in range(end_t)) - abs_end
            all_telems[end_t].text = (all_telems[end_t].text or '')[-remain:] if remain > 0 else ''
        count += 1
        search_start = kw_idx + len(keyword)  # 继续搜索后续关键词
    return count if count > 0 else 0


# ============================================================
#  步骤
# ============================================================

SHORT_WARN_KEYS = {'age', 'sbp', 'dbp'}  # 短值需要警告


def step1_analyze(doc_path):
    print("=" * 60)
    print("  分析模式")
    print("  （正文 + 表格 + 页眉页脚 + 文本框）")
    print("=" * 60)
    print(f"\n文档: {doc_path}\n")

    doc = docx.Document(doc_path)
    results = {}

    for item in ALL_ITEMS:
        val, loc, ctx = find_target(doc, item)
        if val:
            matches = collect_all_matches(doc, val)
            print(f"✅ {item.label}: {val}")
            print(f"   关键词位置: {loc}")
            print(f"   关键词上下文: {ctx}")
            print(f"   全局匹配: {len(matches)} 处")
            for idx, m in enumerate(matches):
                cross = " ⚠️跨run" if m['run_range'][0] != m['run_range'][1] else ""
                print(f"     [{idx}] {m['location']}{cross}")
                if m.get('context'):
                    print(f"          上下文: {m['context']}")
            if item.key in SHORT_WARN_KEYS and len(matches) > 0:
                suspicious = [m for m in matches if m.get('context') and item.label not in m['context'] and '岁' not in (m.get('context') or '')]
                if suspicious:
                    print(f"   ⚠️ 注意：有 {len(suspicious)} 处匹配可能不是 {item.label} 数据")
            print()
            results[item] = {'val': val, 'count': len(matches)}
        else:
            print(f"❌ 未找到 {item.label}\n")

    print("=" * 60)
    print("  分析完成")
    for item, info in results.items():
        print(f"  {item.label}: {info['val']}  (全局 {info['count']} 处)")
    print()
    print("  替换参数 (可任意组合，加 --dry-run 预览):")
    for item in ALL_ITEMS:
        print(f"    --replace-{item.key} 新{item.label}")
    print("=" * 60)


def step2_replace(doc_path, dry_run=False, replacements=None):
    mode = "预览" if dry_run else "替换"
    print("=" * 60)
    print(f"  替换模式 —— {mode}")
    print("=" * 60)

    doc = docx.Document(doc_path)

    if not replacements:
        print("❌ 没有指定替换内容！")
        return

    print(f"\n文档: {doc_path}")
    for label, old, new in replacements:
        print(f"  {label}: '{old}' → '{new}'")
    print()

    total = 0
    age_old = age_new = None
    for label, old, new in replacements:
        if dry_run:
            matches = collect_all_matches(doc, old)
            print(f"📋 [{label}] 将替换 {len(matches)} 处")
            for idx, m in enumerate(matches):
                cross = " ⚠️跨run" if m['run_range'][0] != m['run_range'][1] else ""
                print(f"  [{idx}] {m['location']}{cross}")
            print()
            total += len(matches)
        elif label == '身份证号':
            count = replace_keyword_only(doc, '身份证号', old, new)
            print(f"✅ [{label}] 已替换 {count} 处（仅关键词后）")
            total += count
        else:
            count = replace_all(doc, old, new)
            print(f"✅ [{label}] 已替换 {count} 处")
            total += count
            if label == '年龄' and age_old is None:
                import re
                m_old = re.match(r'(\d+)', old)
                m_new = re.match(r'(\d+)', new)
                if m_old and m_new and len(m_old.group(1)) >= 2:
                    age_old = m_old.group(1)
                    age_new = m_new.group(1)

    # 对 body 做跨 <w:t> 标签替换（处理被 XML 拆散的文本，如日期）
    if not dry_run and replacements:
        ns_t = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'
        ns_p = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'
        for label, old, new in replacements:
            if not old or old == new: continue
            if label == '身份证号': continue  # 身份证号只做关键词精准替换，不跨标签全局替换
            for p_elem in doc._element.body.iter(ns_p):
                t_elems = list(p_elem.iter(ns_t))
                if not t_elems: continue
                joined = ''.join(t.text or '' for t in t_elems)
                if old not in joined: continue
                pos = 0
                while True:
                    idx = joined.find(old, pos)
                    if idx == -1: break
                    length = 0; start_t = 0
                    for i, t in enumerate(t_elems):
                        tl = len(t.text or '')
                        if length + tl > idx: start_t = i; break
                        length += tl
                    else: break
                    off = idx - sum(len(t_elems[j].text or '') for j in range(start_t))
                    end_idx = idx + len(old); end_t = start_t
                    for i in range(start_t, len(t_elems)):
                        if sum(len(t_elems[j].text or '') for j in range(i+1)) >= end_idx: end_t = i; break
                    if start_t == end_t:
                        t_elems[start_t].text = (t_elems[start_t].text or '')[:off] + new + (t_elems[start_t].text or '')[off+len(old):]
                    else:
                        t_elems[start_t].text = (t_elems[start_t].text or '')[:off] + new
                        for i in range(start_t+1, end_t): t_elems[i].text = ''
                        remain = sum(len(t_elems[j].text or '') for j in range(end_t)) - end_idx
                        t_elems[end_t].text = (t_elems[end_t].text or '')[-remain:] if remain > 0 else ''
                    joined = ''.join(t.text or '' for t in t_elems)
                    pos = idx + len(new)
        if age_old and age_new:
            for root in [doc._element.body] + [getattr(s, hn)._element for s in doc.sections for hn in ['header','footer','first_page_header','even_page_header','first_page_footer','even_page_footer'] if getattr(s, hn, None)]:
                for t in root.iter(ns_t):
                    if t.text and t.text.strip() == age_old: t.text = age_new
                    if t.text and ('年龄'+age_old) in (t.text or ''): t.text = t.text.replace('年龄'+age_old, '年龄'+age_new)

    if dry_run:
        print("🔍 以上为预览，文件未被修改。")
    else:
        if total > 0:
            backup = doc_path.replace('.docx', f'_backup_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx')
            shutil.copy(doc_path, backup)
            print(f"\n📦 已备份: {os.path.basename(backup)}")
            doc.save(doc_path)
            print(f"✅ 文件已保存: {doc_path}")
        else:
            print("⚠️ 未找到任何匹配内容，文件未修改。")
    print("=" * 60)


# ============================================================
#  入口
# ============================================================

def main():
    doc_path = "/Users/zhaoshuqing/Desktop/体检报告/郑万军.docx"

    if len(sys.argv) < 2:
        step1_analyze(doc_path)
        return

    dry_run = False
    new_vals = {}

    args = sys.argv[1:]
    i = 0
    while i < len(args):
        if args[i] == '--dry-run':
            dry_run = True; i += 1
        elif args[i] == '--file' and i + 1 < len(args):
            doc_path = args[i + 1]; i += 2
        else:
            matched = False
            for item in ALL_ITEMS:
                flag = f'--replace-{item.key}'
                if args[i] == flag and i + 1 < len(args):
                    new_vals[item.key] = args[i + 1]
                    i += 2; matched = True; break
            if not matched:
                i += 1

    if new_vals:
        doc = docx.Document(doc_path)
        key_to_item = {item.key: item for item in ALL_ITEMS}
        replacements = []

        gender_replacements = []
        for key, new_val in new_vals.items():
            item = key_to_item.get(key)
            if item is None:
                continue
            # 性别特殊处理：使用关键词后精准替换
            if key == 'gender':
                old_val, _, _ = find_target(doc, item)
                if old_val and old_val != new_val:
                    gender_replacements.append((old_val, new_val))
                else:
                    print(f"⚠️ 性别已是 {new_val} 或未找到，跳过")
                continue
            old_val, _, _ = find_target(doc, item)
            if old_val is None:
                print(f"⚠️ 未找到 {item.label}，跳过")
                continue
            # 短数字加单位后缀做精确替换，避免误伤体检编码和检查数据
            if key == 'age':
                replacements.append((item.label, old_val + '岁', new_val + '岁'))
                if len(old_val) >= 2 and len(new_val) >= 2:
                    replacements.append((item.label,
                        old_val[0] + ' ' + old_val[1] + ' ' + '岁',
                        new_val[0] + ' ' + new_val[1] + ' ' + '岁'))
            elif key == 'height':
                # 带cm后缀防止被后续如"60.00"等替换误伤
                replacements.append((item.label, old_val + 'cm', new_val + 'cm'))
            elif key == 'weight':
                # 带Kg后缀防止误伤身高（如160.00中的60.00）
                replacements.append((item.label, old_val + 'Kg', new_val + 'Kg'))
            elif key in ('sbp', 'dbp'):
                replacements.append((item.label, old_val + 'mmHg', new_val + 'mmHg'))
                if key == 'dbp':
                    replacements.append(('心率', old_val + 'bpm', new_val + 'bpm'))
            else:
                replacements.append((item.label, old_val, new_val))

        # 先执行标准替换
        if replacements:
            step2_replace(doc_path, dry_run=dry_run, replacements=replacements)
            # 重新加载文档，因为 step2_replace 已保存
            doc = docx.Document(doc_path)

        # 再执行性别替换
        for old_gen, new_gen in gender_replacements:
            count = replace_gender_in_doc(doc, old_gen, new_gen)
            print(f"✅ [性别] 已替换 {count} 处（{old_gen}→{new_gen}）")
            if not dry_run and count > 0:
                backup = doc_path.replace('.docx', f'_backup_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx')
                shutil.copy(doc_path, backup)
                print(f"📦 已备份: {os.path.basename(backup)}")
                doc.save(doc_path)
                print(f"✅ 文件已保存: {doc_path}")

    else:
        step1_analyze(doc_path)


if __name__ == '__main__':
    main()
